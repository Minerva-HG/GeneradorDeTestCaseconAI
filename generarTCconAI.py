# generarTCconAI.py
# -*- coding: utf-8 -*-
"""
Generador de TestCases (PDF, TXT, CSV, PY, GROOVY, ROBOT, XML, DOCX) con IA (Cohere).
Exportación a TXT/DOCX/XLSX para PDF. La IA usa el documento cargado para generar
casos de prueba (.xlsx) y script Katalon (.groovy).

Requisitos:
    pip install streamlit pdfplumber PyPDF2 python-docx pandas openpyxl cohere

Configuración (variable de entorno o .streamlit/secrets.toml):
    COHERE_API_KEY = tu_clave
    AI_MODEL = command-r-08-2024   (opcional)
"""

# ===== IMPORTS =====
"""Asegura que la salida sea UTF-8 para evitar errores de codificación en Streamlit."""
import sys
import io

if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
if sys.stderr.encoding != 'utf-8':
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

import os
import re
from io import BytesIO
import streamlit as st
import pdfplumber
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
#import xpathGenerator
import importlib

# Extensiones admitidas y tipo de contenido
ALLOWED_EXTENSIONS = ("pdf", "txt", "csv", "py", "groovy", "robot", "xml", "docx")


#--titulo
st.title("Generador de TC y Scripts con AI")


# Botón de Recargar
if st.button("Recargar"):
    st.rerun()

def _get_extension(filename):
    return (filename or "").lower().split(".")[-1] if "." in (filename or "") else ""


# ===== EXTRACCIÓN POR TIPO DE ARCHIVO =====
def extract_text_per_page(file_bytes):
    """Solo para PDF."""
    pages = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            pages.append({"pagina": i, "texto": text.strip()})
    return pages


def extract_metadata_pdf(file_bytes):
    try:
        reader = PdfReader(BytesIO(file_bytes))
        info = reader.metadata or {}
        return {
            "paginas": len(reader.pages),
            "titulo": getattr(info, "title", ""),
            "autor": getattr(info, "author", ""),
        }
    except Exception:
        return {}


def extract_text_plain(file_bytes):
    try:
        return file_bytes.decode("utf-8", errors="replace").strip()
    except Exception:
        return file_bytes.decode("latin-1", errors="replace").strip()


def extract_text_docx(file_bytes):
    try:
        doc = Document(BytesIO(file_bytes))
        return "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    except Exception:
        return ""


def extract_text_csv(file_bytes):
    try:
        text = extract_text_plain(file_bytes)
        df = pd.read_csv(BytesIO(file_bytes), encoding="utf-8", on_bad_lines="skip")
        return df.to_string() if not df.empty else text
    except Exception:
        return extract_text_plain(file_bytes)


def load_document(file_bytes, filename):
    """
    Carga cualquier documento admitido y devuelve (full_text, pages, metadata, is_pdf).
    pages es lista de {pagina, texto} para exportación; is_pdf indica si hay conversión TXT/DOCX/XLSX.
    """
    ext = _get_extension(filename)
    full_text = ""
    pages = []
    metadata = {"tipo": ext, "nombre": filename}

    if ext == "pdf":
        pages = extract_text_per_page(file_bytes)
        full_text = "\n\n".join(p.get("texto", "") for p in pages)
        metadata.update(extract_metadata_pdf(file_bytes))
        return full_text, pages, metadata, True

    if ext == "docx":
        full_text = extract_text_docx(file_bytes)
    elif ext == "csv":
        full_text = extract_text_csv(file_bytes)
    elif ext in ("txt", "py", "groovy", "robot", "xml"):
        full_text = extract_text_plain(file_bytes)
    else:
        full_text = extract_text_plain(file_bytes)

    pages = [{"pagina": 1, "texto": full_text}]
    return full_text, pages, metadata, False


# ===== EXPORTACIONES =====
def to_txt_buffer(pages):
    content = []
    for p in pages:
        content.append("\n=== Página {} ===\n{}".format(p.get("pagina", ""), p.get("texto", "")))
    text = "".join(content)
    return BytesIO(text.encode("utf-8"))


def to_docx_buffer(pages):
    doc = Document()
    doc.add_heading("Documento convertido", level=1)
    for p in pages:
        doc.add_heading("Página {}".format(p.get("pagina", "")), level=2)
        texto = p.get("texto", "")
        if not texto.strip():
            doc.add_paragraph("(Sin texto extraíble)")
        else:
            for par in texto.split("\n\n"):
                doc.add_paragraph(par.strip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def to_xlsx_buffer(pages):
    df = pd.DataFrame(pages)
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Texto")
    buffer.seek(0)
    return buffer


def extract_metadata(file_bytes):
    """Compatibilidad: solo para PDF."""
    return extract_metadata_pdf(file_bytes)


# ===== MENSAJES DE ERROR AMIGABLES =====
def _friendly_ai_error(e: Exception) -> str:
    """Convierte errores de la API Cohere en mensajes entendibles."""
    msg = str(e).lower()
    if "401" in msg or "invalid" in msg or "unauthorized" in msg:
        return "API key de Cohere inválida o expirada. Verifica COHERE_API_KEY en .streamlit/secrets.toml o variables de entorno."
    if "429" in msg or "rate_limit" in msg:
        return "Límite de uso de Cohere alcanzado. Espera un momento o revisa tu cuota en https://dashboard.cohere.com"
    return str(e)


# ===== IA CLIENT (Cohere) =====
class AIClient:
    def __init__(self):
        self.model = os.getenv("AI_MODEL", "command-r-08-2024")
        self.client = None
        self._error = None
        self._init_client()

    def _get_api_key(self, env_name: str, secret_name: str = None) -> str:
        """Obtiene la API key: primero variable de entorno, luego Streamlit secrets."""
        key = os.getenv(env_name, "").strip()
        if not key and secret_name and hasattr(st, "secrets") and st.secrets:
            key = (st.secrets.get(secret_name) or "").strip()
        return key or ""

    def _init_client(self):
        try:
            api_key = self._get_api_key("COHERE_API_KEY", "COHERE_API_KEY")
            if not api_key:
                self._error = "Falta COHERE_API_KEY. Añádela en .streamlit/secrets.toml o variables de entorno."
                return
            import cohere
            self.client = cohere.Client(api_key)
        except Exception as e:
            self.client = None
            self._error = str(e)

    @property
    def ready(self):
        return self.client is not None and self.model

    def summarize(self, text: str, idioma: str = "es") -> str:
        if not self.ready:
            raise RuntimeError("IA no configurada correctamente.")
        system = f"Eres un QA manual experto para hacer casos de pruebas en .xlsx , y eres un QA automation experto que realiza los scripts de pruebas con las mejores practicas de testing {idioma} de forma clara y accionable usando las siguientes herramientas: Selenium, Playwright, Cypress, TestNG, JUnit,katalon studio, robot framework, etc."
        prompt = system + "\n\nTEXTO:\n" + text
        resp = self.client.chat(model=self.model, message=prompt)
        return getattr(resp, "text", "") or getattr(resp, "response", "") or str(resp)

    def answer(self, context: str, question: str, idioma: str = "es") -> str:
        if not self.ready:
            raise RuntimeError("IA no configurada correctamente.")
        system = (
            f"Respondes en {idioma} usando exclusivamente el contexto proporcionado. "
            "Si la información no está, dilo explícitamente."
        )
        truncated = context[:30000]
        prompt = system + "\n\nContexto:\n" + truncated + "\n\nPregunta: " + question
        resp = self.client.chat(model=self.model, message=prompt)
        return getattr(resp, "text", "") or getattr(resp, "response", "") or str(resp)

    def generate_test_cases(self, text: str, idioma: str = "es") -> str:
        """Pide a la IA casos de prueba basados en el documento. Devuelve texto en formato CSV (;) para exportar a Excel."""
        if not self.ready:
            raise RuntimeError("IA no configurada correctamente.")
        system = (
            f"Eres un analista de pruebas. Genera casos de prueba a partir del documento, en {idioma}. "
            "Responde ÚNICAMENTE con una tabla en CSV usando punto y coma (;) como separador. "
            "Primera línea (encabezado): ID;Test Step;Expected;TestData. "
            "Cada fila siguiente es un caso. Usa comillas dobles para campos con ; o saltos de línea. "
            
        )
        truncated = text[:25000]
        prompt = system + "\n\nDOCUMENTO:\n" + truncated
        resp = self.client.chat(model=self.model, message=prompt)
        return getattr(resp, "text", "") or getattr(resp, "response", "") or str(resp)

    def generate_katalon_script(self, text: str, idioma: str = "es") -> str:
        """Pide a la IA un script de pruebas para Katalon Studio en Groovy basado en el documento."""
        if not self.ready:
            raise RuntimeError("IA no configurada correctamente.")
        system = (
            "Eres un experto en automatización con Katalon Studio. Genera un script de prueba en Groovy "
            "basado en el documento (requisitos o especificación). Usa solo keywords de Katalon: "
            "WebUI.openBrowser, WebUI.navigateToUrl, WebUI.click, WebUI.setText, WebUI.verifyElementPresent, "
            "WebUI.verifyTextPresent, WebUI.closeBrowser, etc. Responde ÚNICAMENTE con el código Groovy y todas "
            "las librerías posibles para la automatización de pruebas, "
            "sin explicaciones antes ni después. Incluye comentarios en el código si es útil."
        )
        truncated = text[:25000]
        prompt = system + "\n\nDOCUMENTO:\n" + truncated
        resp = self.client.chat(model=self.model, message=prompt)
        return getattr(resp, "text", "") or getattr(resp, "response", "") or str(resp)


def _parse_test_cases_to_dataframe(raw: str) -> pd.DataFrame:
    """Convierte la respuesta de la IA (CSV con ;) en DataFrame. Limpia bloques markdown si existen."""
    text = raw.strip()
    # Quitar posible bloque ```csv ... ``` o ``` ... ```
    for pattern in (r"```(?:csv)?\s*\n?(.*?)\n?```", r"```\s*\n?(.*?)\n?```"):
        m = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if m:
            text = m.group(1).strip()
    # Leer CSV con ; y manejar posibles errores
    try:
        from io import StringIO
        df = pd.read_csv(StringIO(text), sep=";", encoding="utf-8", quoting=1, on_bad_lines="skip")
    except Exception:
        df = pd.DataFrame(columns=["ID", "Nombre del caso", "Pasos", "Resultado esperado", "Prioridad"])
        # Intentar al menos una fila con el texto
        df.loc[0] = ["TC01", "Caso generado", raw[:200], "Verificar según documento", "Media"]
    return df


def _test_cases_to_xlsx_buffer(df: pd.DataFrame) -> BytesIO:
    """Genera un archivo Excel con los casos de prueba."""
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Casos de prueba")
    buffer.seek(0)
    return buffer


# ===== STREAMLIT UI =====
st.set_page_config(page_title="Generador de TestCases", page_icon="🤖")



st.caption("PDF, TXT, CSV, PY, GROOVY, ROBOT, XML, DOCX — IA genera casos de prueba y script Katalon")

uploaded_file = st.file_uploader(
    "Sube un documento",
    type=list(ALLOWED_EXTENSIONS),
    help="El contenido se usará para que la IA genere casos de prueba (.xlsx) y script Katalon (.groovy).",
)

if uploaded_file:
    file_bytes = uploaded_file.read()
    full_text, pages, metadata, is_pdf = load_document(file_bytes, uploaded_file.name)
    import unicodedata
    def _sanitize_filename(name: str) -> str:
        # Normaliza y convierte caracteres con tilde a su equivalente ASCII (ej: é -> e, ñ -> n)
        nfkd = unicodedata.normalize("NFKD", name)
        ascii_name = nfkd.encode("ascii", errors="ignore").decode("ascii")
        # Elimina caracteres inválidos en nombres de archivo
        return re.sub(r'[<>:"/\\|?*]', "_", ascii_name).strip() or "documento"
    _base_name = _sanitize_filename(os.path.splitext(uploaded_file.name or "documento")[0].strip())

    st.subheader("Información del documento")
    st.json(metadata)

    if is_pdf:
        st.subheader("Descargar archivo convertido")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.download_button("⬇ TXT", to_txt_buffer(pages), file_name=f"{_base_name}.txt", mime="text/plain")
        with col2:
            st.download_button(
                "⬇ DOCX", to_docx_buffer(pages),
                file_name=f"{_base_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col3:
            st.download_button(
                "⬇ XLSX", to_xlsx_buffer(pages),
                file_name=f"{_base_name}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        st.markdown("---")

    st.subheader("🧠 IA (opcional)")
    if st.checkbox("Activar IA"):
        ai = AIClient()
        if not ai.ready:
            msg = getattr(ai, "_error", None) or "Configura COHERE_API_KEY."
            st.warning(f"IA no configurada: {msg}")
            st.code('COHERE_API_KEY = "tu_clave"\nAI_MODEL = "command-r-08-2024"  # opcional', language="toml")
        else:
            if not full_text.strip():
                st.error("No se pudo extraer texto del documento.")
            else:
                st.markdown("**Generar casos de prueba y script Katalon** (basados en el documento cargado)")
                col_cp, col_kt = st.columns(2)
                with col_cp:
                    if st.button("📋 Generar casos de prueba (XLSX)", key="btn_casos"):
                        try:
                            with st.spinner("Generando casos de prueba…"):
                                raw = ai.generate_test_cases(full_text, idioma="es")
                                df = _parse_test_cases_to_dataframe(raw)
                                buf = _test_cases_to_xlsx_buffer(df)
                                st.session_state["casos_prueba_xlsx"] = buf.getvalue()
                            st.rerun()
                        except Exception as e:
                            st.error(_friendly_ai_error(e))
                    if st.session_state.get("casos_prueba_xlsx"):
                        st.download_button(
                            "⬇ Descargar casos de prueba (.xlsx)",
                            data=st.session_state["casos_prueba_xlsx"],
                            file_name=f"{_base_name}_casos_de_prueba.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key="dl_casos_prueba",
                        )
                with col_kt:
                    if st.button("🔧 Generar script Katalon (.groovy)", key="btn_katalon"):
                        try:
                            with st.spinner("Generando script Katalon…"):
                                raw = ai.generate_katalon_script(full_text, idioma="es")
                                code = raw.strip()
                                m = re.search(r"```(?:groovy)?\s*\n?(.*?)\n?```", code, re.DOTALL | re.IGNORECASE)
                                if m:
                                    code = m.group(1).strip()
                                st.session_state["katalon_groovy"] = code
                            st.rerun()
                        except Exception as e:
                            st.error(_friendly_ai_error(e))
                    if st.session_state.get("katalon_groovy"):
                        st.download_button(
                            "⬇ Descargar script Katalon (.groovy)",
                            data=st.session_state["katalon_groovy"].encode("utf-8"),
                            file_name=f"{_base_name}.groovy",
                            mime="text/plain",
                            key="dl_katalon",
                        )

                st.markdown("---")
                st.markdown("**Preguntas sobre el documento**")
                q = st.text_input("Escribe tu pregunta sobre el documento")
                if st.button("🔍 Responder") and q.strip():
                    try:
                        with st.spinner("Buscando respuesta…"):
                            respuesta = ai.answer(full_text, q, idioma="es")
                        st.markdown("### Respuesta")
                        st.write(respuesta)
                    except Exception as e:
                        st.error(_friendly_ai_error(e))
else:
    st.info("Sube un documento (PDF, TXT, CSV, PY, GROOVY, ROBOT, XML o DOCX) para comenzar")


#----BOTON DE XPATHOR---
#st.title("Xpathor")

#st.markdown("[Abrir XpathGenerator](http://localhost:8502)")
    
