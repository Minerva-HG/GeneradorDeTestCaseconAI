# generarTCconAI.py
# -*- coding: utf-8 -*-
"""
Generador de TestCases (PDF, TXT, CSV, PY, GROOVY, ROBOT, XML, DOCX) con IA (Cohere).
Exportación a TXT/DOCX/XLSX para PDF. La IA usa el documento cargado para generar
casos de prueba (.xlsx) y script Katalon (.groovy).

Requisitos:
    pip install streamlit pdfplumber PyPDF2 python-docx pandas openpyxl cohere

Configuración (variable de entorno o .streamlit/secrets.toml):
    COHERE_API_KEY = tu_clave
    AI_MODEL = command-r-08-2024   (opcional)
"""

# ===== FORZAR UTF-8 EN WINDOWS =====
import sys
import io
if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if sys.stderr.encoding and sys.stderr.encoding.lower() != "utf-8":
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

# ===== IMPORTS =====
import os
import re
import unicodedata
from io import BytesIO, StringIO

import streamlit as st
import pdfplumber
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd

# ===== CONFIGURACIÓN DE PÁGINA (debe ir primero) =====
st.set_page_config(page_title="Generador de TestCases", page_icon="🤖")

# Extensiones admitidas
ALLOWED_EXTENSIONS = ("pdf", "txt", "csv", "py", "groovy", "robot", "xml", "docx")


# ===== UTILIDADES =====
def _sanitize_filename(name: str) -> str:
    """
    Convierte un nombre de archivo a ASCII seguro:
    - Normaliza caracteres con tilde (é->e, ñ->n, etc.)
    - Elimina caracteres inválidos en nombres de archivo
    """
    nfkd = unicodedata.normalize("NFKD", name)
    ascii_name = nfkd.encode("ascii", errors="ignore").decode("ascii")
    return re.sub(r'[<>:"/\\|?*\s]+', "_", ascii_name).strip("_") or "documento"


def _get_extension(filename):
    return (filename or "").lower().split(".")[-1] if "." in (filename or "") else ""


# ===== EXTRACCIÓN POR TIPO DE ARCHIVO =====
def extract_text_per_page(file_bytes):
    """Solo para PDF."""
    pages = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            pages.append({"pagina": i, "texto": text.strip()})
    return pages


def extract_metadata_pdf(file_bytes):
    try:
        reader = PdfReader(BytesIO(file_bytes))
        info = reader.metadata or {}
        return {
            "paginas": len(reader.pages),
            "titulo": getattr(info, "title", "") or "",
            "autor": getattr(info, "author", "") or "",
        }
    except Exception:
        return {}


def extract_text_plain(file_bytes):
    try:
        return file_bytes.decode("utf-8", errors="replace").strip()
    except Exception:
        return file_bytes.decode("latin-1", errors="replace").strip()


def extract_text_docx(file_bytes):
    try:
        doc = Document(BytesIO(file_bytes))
        return "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    except Exception:
        return ""


def extract_text_csv(file_bytes):
    try:
        df = pd.read_csv(BytesIO(file_bytes), encoding="utf-8", on_bad_lines="skip")
        return df.to_string() if not df.empty else extract_text_plain(file_bytes)
    except Exception:
        return extract_text_plain(file_bytes)


def load_document(file_bytes, filename):
    """
    Carga cualquier documento admitido y devuelve (full_text, pages, metadata, is_pdf).
    pages es lista de {pagina, texto} para exportación.
    is_pdf indica si el documento original es PDF (y se ofrecen conversiones).
    """
    ext = _get_extension(filename)
    metadata = {"tipo": ext, "nombre": filename}

    if ext == "pdf":
        pages = extract_text_per_page(file_bytes)
        full_text = "\n\n".join(p.get("texto", "") for p in pages)
        metadata.update(extract_metadata_pdf(file_bytes))
        return full_text, pages, metadata, True

    extractors = {
        "docx": extract_text_docx,
        "csv": extract_text_csv,
    }
    full_text = extractors.get(ext, extract_text_plain)(file_bytes)
    pages = [{"pagina": 1, "texto": full_text}]
    return full_text, pages, metadata, False


# ===== EXPORTACIONES =====
def to_txt_buffer(pages):
    content = "\n".join(
        "\n=== Página {} ===\n{}".format(p.get("pagina", ""), p.get("texto", ""))
        for p in pages
    )
    return BytesIO(content.encode("utf-8"))


def to_docx_buffer(pages):
    doc = Document()
    doc.add_heading("Documento convertido", level=1)
    for p in pages:
        doc.add_heading("Página {}".format(p.get("pagina", "")), level=2)
        texto = p.get("texto", "")
        if not texto.strip():
            doc.add_paragraph("(Sin texto extraíble)")
        else:
            for par in texto.split("\n\n"):
                if par.strip():
                    doc.add_paragraph(par.strip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def to_xlsx_buffer(pages):
    df = pd.DataFrame(pages)
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Texto")
    buffer.seek(0)
    return buffer


# ===== MENSAJES DE ERROR AMIGABLES =====
def _friendly_ai_error(e: Exception) -> str:
    """Convierte errores de la API Cohere en mensajes entendibles."""
    msg = str(e).lower()
    if "401" in msg or "invalid" in msg or "unauthorized" in msg:
        return "API key de Cohere inválida o expirada. Verifica COHERE_API_KEY en .streamlit/secrets.toml o variables de entorno."
    if "429" in msg or "rate_limit" in msg:
        return "Límite de uso de Cohere alcanzado. Espera un momento o revisa tu cuota en https://dashboard.cohere.com"
    if "timeout" in msg or "timed out" in msg:
        return "La solicitud a Cohere tardó demasiado. Intenta con un documento más corto o vuelve a intentarlo."
    if "connection" in msg or "network" in msg:
        return "Error de conexión con Cohere. Verifica tu conexión a internet."
    return str(e)


# ===== IA CLIENT (Cohere) =====
@st.cache_resource
def _build_ai_client():
    """Construye y cachea el cliente de Cohere para evitar reinicializarlo en cada rerun."""
    return AIClient()


class AIClient:
    def __init__(self):
        self.model = os.getenv("AI_MODEL", "command-r-08-2024")
        self.client = None
        self._error = None
        self._init_client()

    def _get_api_key(self, env_name: str, secret_name: str = None) -> str:
        """Obtiene la API key: primero variable de entorno, luego Streamlit secrets."""
        key = os.getenv(env_name, "").strip()
        if not key and secret_name and hasattr(st, "secrets") and st.secrets:
            key = (st.secrets.get(secret_name) or "").strip()
        return key or ""

    def _init_client(self):
        try:
            api_key = self._get_api_key("COHERE_API_KEY", "COHERE_API_KEY")
            if not api_key:
                self._error = "Falta COHERE_API_KEY. Añádela en .streamlit/secrets.toml o variables de entorno."
                return
            import cohere
            self.client = cohere.Client(api_key)
        except Exception as e:
            self.client = None
            self._error = str(e)

    @property
    def ready(self):
        return self.client is not None and bool(self.model)

    def _chat(self, prompt: str) -> str:
        """Método interno común para llamar a la API."""
        resp = self.client.chat(model=self.model, message=prompt)
        return getattr(resp, "text", "") or getattr(resp, "response", "") or str(resp)

    def answer(self, context: str, question: str, idioma: str = "es") -> str:
        if not self.ready:
            raise RuntimeError("IA no configurada correctamente.")
        system = (
            f"Respondes en {idioma} usando exclusivamente el contexto proporcionado. "
            "Si la información no está en el contexto, dilo explícitamente."
        )
        prompt = system + "\n\nContexto:\n" + context[:30000] + "\n\nPregunta: " + question
        return self._chat(prompt)

    def generate_test_cases(self, text: str, idioma: str = "es") -> str:
        """Genera casos de prueba en CSV (;) a partir del documento."""
        if not self.ready:
            raise RuntimeError("IA no configurada correctamente.")
        system = (
            f"Eres un analista de pruebas QA experto. Genera casos de prueba a partir del documento, en {idioma}. "
            "Responde ÚNICAMENTE con una tabla en CSV usando punto y coma (;) como separador. "
            "Primera línea (encabezado): ID;Test Step;Expected;TestData. "
            "Cada fila siguiente es un caso de prueba. Usa comillas dobles para campos que contengan ; o saltos de línea."
        )
        prompt = system + "\n\nDOCUMENTO:\n" + text[:25000]
        return self._chat(prompt)

    def generate_katalon_script(self, text: str, idioma: str = "es") -> str:
        """Genera un script Groovy para Katalon Studio basado en el documento."""
        if not self.ready:
            raise RuntimeError("IA no configurada correctamente.")
        system = (
            "Eres un experto en automatización con Katalon Studio. Genera un script de prueba en Groovy "
            "basado en el documento (requisitos o especificación). Usa keywords de Katalon como: "
            "WebUI.openBrowser, WebUI.navigateToUrl, WebUI.click, WebUI.setText, WebUI.verifyElementPresent, "
            "WebUI.verifyTextPresent, WebUI.closeBrowser, etc. "
            "Responde ÚNICAMENTE con el código Groovy, sin explicaciones antes ni después. "
            "Incluye comentarios en el código donde sea útil."
        )
        prompt = system + "\n\nDOCUMENTO:\n" + text[:25000]
        return self._chat(prompt)


# ===== PARSEO DE CASOS DE PRUEBA =====
def _parse_test_cases_to_dataframe(raw: str) -> pd.DataFrame:
    """Convierte la respuesta de la IA (CSV con ;) en DataFrame. Limpia bloques markdown si existen."""
    text = raw.strip()
    # Quitar posible bloque ```csv ... ``` o ``` ... ```
    m = re.search(r"```(?:csv)?\s*\n?(.*?)\n?```", text, re.DOTALL | re.IGNORECASE)
    if m:
        text = m.group(1).strip()
    try:
        df = pd.read_csv(StringIO(text), sep=";", encoding="utf-8", quoting=1, on_bad_lines="skip")
        # Normalizar nombres de columna eliminando espacios extra
        df.columns = [c.strip() for c in df.columns]
    except Exception:
        df = pd.DataFrame(columns=["ID", "Test Step", "Expected", "TestData"])
        df.loc[0] = ["TC01", "Caso generado automáticamente", raw[:200], "Verificar según documento"]
    return df


def _test_cases_to_xlsx_buffer(df: pd.DataFrame) -> BytesIO:
    """Genera un archivo Excel con los casos de prueba."""
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Casos de prueba")
    buffer.seek(0)
    return buffer


def _clean_groovy_code(raw: str) -> str:
    """Extrae el código Groovy limpio de la respuesta de la IA."""
    code = raw.strip()
    m = re.search(r"```(?:groovy)?\s*\n?(.*?)\n?```", code, re.DOTALL | re.IGNORECASE)
    return m.group(1).strip() if m else code


# ===== STREAMLIT UI =====
st.title("Generador de TC y Scripts con AI")
st.caption("PDF, TXT, CSV, PY, GROOVY, ROBOT, XML, DOCX — IA genera casos de prueba y script Katalon")

if st.button("Recargar"):
    st.rerun()

uploaded_file = st.file_uploader(
    "Sube un documento",
    type=list(ALLOWED_EXTENSIONS),
    help="El contenido se usará para que la IA genere casos de prueba (.xlsx) y script Katalon (.groovy).",
)

if uploaded_file:
    file_bytes = uploaded_file.read()
    full_text, pages, metadata, is_pdf = load_document(file_bytes, uploaded_file.name)
    _base_name = _sanitize_filename(os.path.splitext(uploaded_file.name or "documento")[0].strip())

    st.subheader("Información del documento")
    st.json(metadata)

    # Mostrar estadísticas básicas del texto extraído
    word_count = len(full_text.split()) if full_text.strip() else 0
    st.caption(f"Texto extraído: {word_count:,} palabras · {len(full_text):,} caracteres")

    if is_pdf:
        st.subheader("Descargar archivo convertido")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.download_button(
                "⬇ TXT", to_txt_buffer(pages),
                file_name=f"{_base_name}.txt",
                mime="text/plain",
            )
        with col2:
            st.download_button(
                "⬇ DOCX", to_docx_buffer(pages),
                file_name=f"{_base_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col3:
            st.download_button(
                "⬇ XLSX", to_xlsx_buffer(pages),
                file_name=f"{_base_name}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        st.markdown("---")

    st.subheader("🧠 IA (opcional)")
    if st.checkbox("Activar IA"):
        ai = _build_ai_client()
        if not ai.ready:
            msg = getattr(ai, "_error", None) or "Configura COHERE_API_KEY."
            st.warning(f"IA no configurada: {msg}")
            st.code('COHERE_API_KEY = "tu_clave"\nAI_MODEL = "command-r-08-2024"  # opcional', language="toml")
        else:
            if not full_text.strip():
                st.error("No se pudo extraer texto del documento.")
            else:
                st.markdown("**Generar casos de prueba y script Katalon** (basados en el documento cargado)")
                col_cp, col_kt = st.columns(2)

                # --- Columna: Casos de prueba ---
                with col_cp:
                    if st.button("📋 Generar casos de prueba (XLSX)", key="btn_casos"):
                        try:
                            with st.spinner("Generando casos de prueba…"):
                                raw = ai.generate_test_cases(full_text, idioma="es")
                                df = _parse_test_cases_to_dataframe(raw)
                                buf = _test_cases_to_xlsx_buffer(df)
                                st.session_state["casos_prueba_xlsx"] = buf.getvalue()
                                st.session_state["casos_prueba_df"] = df
                            st.rerun()
                        except Exception as e:
                            st.error(_friendly_ai_error(e))

                    if st.session_state.get("casos_prueba_xlsx"):
                        st.download_button(
                            "⬇ Descargar casos de prueba (.xlsx)",
                            data=st.session_state["casos_prueba_xlsx"],
                            file_name=f"{_base_name}_casos_de_prueba.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key="dl_casos_prueba",
                        )
                        # Preview de los casos generados
                        if "casos_prueba_df" in st.session_state:
                            with st.expander("Vista previa de casos de prueba"):
                                st.dataframe(st.session_state["casos_prueba_df"], use_container_width=True)

                # --- Columna: Script Katalon ---
                with col_kt:
                    if st.button("🔧 Generar script Katalon (.groovy)", key="btn_katalon"):
                        try:
                            with st.spinner("Generando script Katalon…"):
                                raw = ai.generate_katalon_script(full_text, idioma="es")
                                code = _clean_groovy_code(raw)
                                st.session_state["katalon_groovy"] = code
                            st.rerun()
                        except Exception as e:
                            st.error(_friendly_ai_error(e))

                    if st.session_state.get("katalon_groovy"):
                        st.download_button(
                            "⬇ Descargar script Katalon (.groovy)",
                            data=st.session_state["katalon_groovy"].encode("utf-8"),
                            file_name=f"{_base_name}.groovy",
                            mime="text/plain",
                            key="dl_katalon",
                        )
                        # Preview del script con syntax highlighting
                        with st.expander("Vista previa del script Katalon"):
                            st.code(st.session_state["katalon_groovy"], language="groovy")

                st.markdown("---")

                # --- Preguntas sobre el documento ---
                st.markdown("**Preguntas sobre el documento**")
                q = st.text_input("Escribe tu pregunta sobre el documento")
                if st.button("🔍 Responder") and q.strip():
                    try:
                        with st.spinner("Buscando respuesta…"):
                            respuesta = ai.answer(full_text, q, idioma="es")
                        st.markdown("### Respuesta")
                        st.write(respuesta)
                    except Exception as e:
                        st.error(_friendly_ai_error(e))
else:
    st.info("Sube un documento (PDF, TXT, CSV, PY, GROOVY, ROBOT, XML o DOCX) para comenzar.")

st.markdown("---")
st.markdown("[Abrir XpathGenerator](http://localhost:8502)")
